#!/usr/bin/env python
"""
Generate policy documents from templates.
This script takes a Word template, replaces placeholders, and inserts a logo.
"""

import os
import sys
import traceback
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Simple function to write to log file
def log(message):
    with open(os.path.join(os.path.dirname(__file__), "generate_policy.log"), "a") as f:
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        f.write(f"{timestamp} - {message}\n")

def replace_text_in_document(doc, placeholder, replacement):
    """Replace all instances of a placeholder in all paragraphs and tables."""
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            log(f"Found placeholder {placeholder} in paragraph")
            for run in paragraph.runs:
                run.text = run.text.replace(placeholder, replacement)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        log(f"Found placeholder {placeholder} in table cell")
                        for run in paragraph.runs:
                            run.text = run.text.replace(placeholder, replacement)

def insert_image(doc, placeholder, image_path):
    """Find placeholder text and replace it with an image."""
    if not os.path.exists(image_path):
        log(f"Image file not found: {image_path}")
        return False
    
    # Try to find paragraphs with "logo" text
    image_found = False
    for i, paragraph in enumerate(doc.paragraphs):
        para_text = paragraph.text.strip()
        if "logo" in para_text.lower() or placeholder in para_text:
            log(f"Found image placeholder in paragraph {i}")
            
            # Clear the paragraph text
            for run in paragraph.runs:
                run.text = ""
            
            # Add the image
            try:
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(2))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                image_found = True
                log(f"Successfully inserted image at paragraph {i}")
                break
            except Exception as e:
                log(f"Error inserting image: {str(e)}")
    
    # If no placeholder found, try the first paragraph
    if not image_found:
        try:
            if doc.paragraphs:
                paragraph = doc.paragraphs[0]
                # Clear the paragraph text
                for run in paragraph.runs:
                    run.text = ""
                
                # Add the image
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(2))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                image_found = True
                log("Successfully added logo to the first paragraph")
        except Exception as e:
            log(f"Error adding logo to first paragraph: {str(e)}")
    
    return image_found

def main():
    try:
        # Get command line arguments
        if len(sys.argv) < 5:
            log("Error: Insufficient arguments")
            log("Usage: python generate_policy.py <template_path> <output_dir> <company_name> <owner_name> [logo_path]")
            sys.exit(1)
        
        template_path = sys.argv[1]
        output_dir = sys.argv[2]
        company_name = sys.argv[3]
        owner_name = sys.argv[4]
        logo_path = sys.argv[5] if len(sys.argv) > 5 and sys.argv[5] != "null" else None
        
        log("Starting document generation")
        log(f"Template: {template_path}")
        log(f"Output dir: {output_dir}")
        log(f"Company: {company_name}")
        log(f"Owner: {owner_name}")
        log(f"Logo: {logo_path or 'None'}")
        
        # Verify the template path
        if not os.path.isfile(template_path):
            log(f"Template file does not exist: {template_path}")
            sys.exit(1)
        
        # Create output directory
        os.makedirs(output_dir, exist_ok=True)
        
        # Load template
        try:
            doc = Document(template_path)
            log(f"Successfully loaded template: {template_path}")
        except Exception as e:
            log(f"Error loading template: {str(e)}")
            sys.exit(1)
        
        # Replace placeholders
        replace_text_in_document(doc, "{Company_Name}", company_name)
        replace_text_in_document(doc, "{Owner_Name}", owner_name)
        replace_text_in_document(doc, "{current_date}", datetime.now().strftime("%m/%d/%Y"))
        
        # Insert logo if provided
        if logo_path and os.path.exists(logo_path):
            insert_image(doc, "{Company_Logo}", logo_path)
        
        # Generate output filename
        template_name = os.path.basename(template_path).replace('.docx', '')
        sanitized_company_name = company_name.replace(' ', '_')
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{sanitized_company_name}_{template_name}_{timestamp}.docx"
        output_path = os.path.join(output_dir, output_filename)
        
        # Save document
        try:
            doc.save(output_path)
            log(f"Document saved successfully: {output_path}")
            
            # Write output path to file
            output_file = os.path.join(os.path.dirname(__file__), "last_output_path.txt")
            with open(output_file, "w") as f:
                f.write(output_path)
            
            # Also print path to stdout
            print(output_path)
            return 0
        except Exception as e:
            log(f"Error saving document: {str(e)}")
            sys.exit(1)
    
    except Exception as e:
        log(f"Unhandled exception: {str(e)}")
        log(traceback.format_exc())
        sys.exit(1)

if __name__ == "__main__":
    sys.exit(main()) 